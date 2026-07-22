"""
ResumeAgent — generates a tailored, design-varied HTML resume.

Flow:
  1. Optional: Claude Vision analyses a reference image to extract style cues.
  2. Gap analysis and HTML generation run in parallel.
     - Gap analysis → missing_data_requests (questions for the user).
     - HTML generation → complete, ATS-safe, self-contained HTML document.
  3. On regeneration the same flow runs with supplemental_answers folded in.
"""
from __future__ import annotations

import asyncio
import base64
import io
import json
import logging
import random
from typing import Optional

import anthropic

from backend.services.llm_client import call_llm
from backend.services.user_profile import get_profile
from models.job import JobMatch

logger = logging.getLogger(__name__)

_MODEL      = "claude-sonnet-4-6"
_MAX_TOKENS = 8000   # HTML resumes are large


# ── Design attribute pools ─────────────────────────────────────────────────────
# Each axis is sampled independently. 7×9×9×6 = 3,402 base combinations;
# Claude's creative interpretation within each spec multiplies that further.

_STRUCTURES = [
    {
        "label": "Left Sidebar 30/70",
        "description": (
            "CSS grid with two columns: fixed 260px left sidebar and remaining width for main content. "
            "Sidebar holds: candidate name + current title at the very top, contact placeholders, "
            "Skills section, Education section. "
            "Main content holds: Professional Summary, then Experience entries, then Volunteering. "
            "Sidebar and main scroll together; sidebar has its own background colour."
        ),
    },
    {
        "label": "Right Sidebar 70/30",
        "description": (
            "CSS grid with two columns: main content at left taking 70% of width, "
            "then a 240px right sidebar. "
            "Header (name + title + contact) spans the full width above both columns. "
            "Main column: Summary, Experience. "
            "Right sidebar: Skills (as tags), Education, Volunteering."
        ),
    },
    {
        "label": "Single Column",
        "description": (
            "Single centered column, max-width 750px, 40px horizontal padding, 48px vertical padding. "
            "Name and title left-aligned in a header block at the top. "
            "Sections flow vertically: Summary · Experience · Education · Skills · Volunteering. "
            "Whitespace and typography carry the visual weight — no sidebars."
        ),
    },
    {
        "label": "Header Band",
        "description": (
            "Full-width solid header band spanning 100% of page width. "
            "Header band contains: name, current title, and contact line, all white on dark background. "
            "Below the band: two CSS columns — 65% left for Experience + Summary, "
            "35% right for Skills + Education + Volunteering. "
            "Right column has a subtle tinted background."
        ),
    },
    {
        "label": "Asymmetric Grid",
        "description": (
            "CSS grid with an asymmetric 3fr/2fr (60%/40%) split below a full-width header block. "
            "Header block: name large, title, contact — left aligned, no background colour. "
            "Left column: Experience (primary), Summary. "
            "Right column: Skills as tags, Education, Volunteering, Contact repeated compact."
        ),
    },
    {
        "label": "Centered Nameplate",
        "description": (
            "Single column with a centered nameplate header: name very large and centered, "
            "title centered below, two thin horizontal rules framing the header block. "
            "All section content below is left-aligned. "
            "Sections: Experience · Skills · Education · Summary · Volunteering."
        ),
    },
    {
        "label": "Wide Left + Thin Right",
        "description": (
            "CSS grid: 1fr left column for Experience + Summary and a narrow 200px right column "
            "for Skills + Contact + Education. "
            "Name and title sit above the grid in a left-aligned header row. "
            "Right column items are compact, tight line-height, no headings — just labelled lists."
        ),
    },
]

_TYPOGRAPHY = [
    {
        "label": "Playfair Display / Lato",
        "import": "@import url('https://fonts.googleapis.com/css2?family=Playfair+Display:wght@400;600;700&family=Lato:wght@300;400;700&display=swap');",
        "heading_font": "'Playfair Display', Georgia, serif",
        "body_font":    "'Lato', 'Helvetica Neue', sans-serif",
        "name_size":    "28px",
        "heading_size": "10px",
        "heading_weight": "600",
        "heading_transform": "uppercase",
        "heading_spacing":   "2px",
        "body_size":    "10px",
    },
    {
        "label": "Roboto / Merriweather",
        "import": "@import url('https://fonts.googleapis.com/css2?family=Roboto:wght@300;400;500;700&family=Merriweather:wght@300;400;700&display=swap');",
        "heading_font": "'Roboto', Arial, sans-serif",
        "body_font":    "'Merriweather', Georgia, serif",
        "name_size":    "26px",
        "heading_size": "10.5px",
        "heading_weight": "700",
        "heading_transform": "uppercase",
        "heading_spacing":   "1.5px",
        "body_size":    "9.5px",
    },
    {
        "label": "Inter / Inter",
        "import": "@import url('https://fonts.googleapis.com/css2?family=Inter:wght@300;400;500;600;700&display=swap');",
        "heading_font": "'Inter', system-ui, sans-serif",
        "body_font":    "'Inter', system-ui, sans-serif",
        "name_size":    "24px",
        "heading_size": "10px",
        "heading_weight": "600",
        "heading_transform": "uppercase",
        "heading_spacing":   "2.5px",
        "body_size":    "10px",
    },
    {
        "label": "Space Mono / Work Sans",
        "import": "@import url('https://fonts.googleapis.com/css2?family=Space+Mono:wght@400;700&family=Work+Sans:wght@300;400;500;600&display=swap');",
        "heading_font": "'Space Mono', 'Courier New', monospace",
        "body_font":    "'Work Sans', 'Helvetica Neue', sans-serif",
        "name_size":    "22px",
        "heading_size": "9.5px",
        "heading_weight": "700",
        "heading_transform": "uppercase",
        "heading_spacing":   "3px",
        "body_size":    "10px",
    },
    {
        "label": "Cormorant Garamond / Source Sans Pro",
        "import": "@import url('https://fonts.googleapis.com/css2?family=Cormorant+Garamond:wght@400;500;600;700&family=Source+Sans+Pro:wght@300;400;600&display=swap');",
        "heading_font": "'Cormorant Garamond', Garamond, serif",
        "body_font":    "'Source Sans Pro', 'Helvetica Neue', sans-serif",
        "name_size":    "32px",
        "heading_size": "11px",
        "heading_weight": "600",
        "heading_transform": "small-caps",
        "heading_spacing":   "1px",
        "body_size":    "10px",
    },
    {
        "label": "Raleway / Crimson Text",
        "import": "@import url('https://fonts.googleapis.com/css2?family=Raleway:wght@300;400;500;600;700&family=Crimson+Text:ital,wght@0,400;0,600;1,400&display=swap');",
        "heading_font": "'Raleway', 'Helvetica Neue', sans-serif",
        "body_font":    "'Crimson Text', Georgia, serif",
        "name_size":    "28px",
        "heading_size": "10px",
        "heading_weight": "600",
        "heading_transform": "uppercase",
        "heading_spacing":   "2px",
        "body_size":    "11px",
    },
    {
        "label": "Josefin Sans / Libre Baskerville",
        "import": "@import url('https://fonts.googleapis.com/css2?family=Josefin+Sans:wght@300;400;600;700&family=Libre+Baskerville:ital,wght@0,400;0,700;1,400&display=swap');",
        "heading_font": "'Josefin Sans', 'Helvetica Neue', sans-serif",
        "body_font":    "'Libre Baskerville', Georgia, serif",
        "name_size":    "26px",
        "heading_size": "10.5px",
        "heading_weight": "600",
        "heading_transform": "uppercase",
        "heading_spacing":   "3px",
        "body_size":    "9.5px",
    },
    {
        "label": "Nunito / Noto Serif",
        "import": "@import url('https://fonts.googleapis.com/css2?family=Nunito:wght@300;400;600;700&family=Noto+Serif:wght@400;700&display=swap');",
        "heading_font": "'Nunito', 'Helvetica Neue', sans-serif",
        "body_font":    "'Noto Serif', Georgia, serif",
        "name_size":    "26px",
        "heading_size": "10px",
        "heading_weight": "700",
        "heading_transform": "uppercase",
        "heading_spacing":   "1.5px",
        "body_size":    "10px",
    },
    {
        "label": "DM Serif / DM Sans",
        "import": "@import url('https://fonts.googleapis.com/css2?family=DM+Serif+Display&family=DM+Sans:wght@300;400;500;700&display=swap');",
        "heading_font": "'DM Serif Display', Georgia, serif",
        "body_font":    "'DM Sans', system-ui, sans-serif",
        "name_size":    "30px",
        "heading_size": "10px",
        "heading_weight": "400",
        "heading_transform": "uppercase",
        "heading_spacing":   "2px",
        "body_size":    "10px",
    },
]

_PALETTES = [
    {
        "label": "Slate & Indigo",
        "page_bg":      "#F8FAFC",
        "sidebar_bg":   "#F1F5F9",
        "main_bg":      "#FFFFFF",
        "text_primary": "#1E293B",
        "text_muted":   "#64748B",
        "accent":       "#4F46E5",
        "accent_light": "#EEF2FF",
        "header_bg":    "#1E293B",
        "header_text":  "#FFFFFF",
        "sidebar_text": "#1E293B",
        "sidebar_muted":"#64748B",
        "divider":      "#E2E8F0",
    },
    {
        "label": "Charcoal & Mint",
        "page_bg":      "#F9FAFB",
        "sidebar_bg":   "#F0FDF4",
        "main_bg":      "#FFFFFF",
        "text_primary": "#111827",
        "text_muted":   "#6B7280",
        "accent":       "#059669",
        "accent_light": "#ECFDF5",
        "header_bg":    "#111827",
        "header_text":  "#FFFFFF",
        "sidebar_text": "#111827",
        "sidebar_muted":"#6B7280",
        "divider":      "#D1FAE5",
    },
    {
        "label": "Navy & Gold",
        "page_bg":      "#F8F7F4",
        "sidebar_bg":   "#1E3A5F",
        "main_bg":      "#FFFFFF",
        "text_primary": "#1A2744",
        "text_muted":   "#5C6E8A",
        "accent":       "#B8972A",
        "accent_light": "#FEF9EC",
        "header_bg":    "#1A2744",
        "header_text":  "#FFFFFF",
        "sidebar_text": "#FFFFFF",
        "sidebar_muted":"#94A3B8",
        "divider":      "#E5E0D0",
    },
    {
        "label": "Monochromatic Dark",
        "page_bg":      "#0F172A",
        "sidebar_bg":   "#1E293B",
        "main_bg":      "#1E293B",
        "text_primary": "#F1F5F9",
        "text_muted":   "#94A3B8",
        "accent":       "#38BDF8",
        "accent_light": "#0C4A6E",
        "header_bg":    "#0F172A",
        "header_text":  "#F1F5F9",
        "sidebar_text": "#F1F5F9",
        "sidebar_muted":"#94A3B8",
        "divider":      "#334155",
    },
    {
        "label": "Crisp White & Emerald",
        "page_bg":      "#FFFFFF",
        "sidebar_bg":   "#ECFDF5",
        "main_bg":      "#FFFFFF",
        "text_primary": "#064E3B",
        "text_muted":   "#6B7280",
        "accent":       "#10B981",
        "accent_light": "#D1FAE5",
        "header_bg":    "#064E3B",
        "header_text":  "#FFFFFF",
        "sidebar_text": "#064E3B",
        "sidebar_muted":"#6B7280",
        "divider":      "#A7F3D0",
    },
    {
        "label": "Warm Sand & Terracotta",
        "page_bg":      "#FDF6EC",
        "sidebar_bg":   "#F5EDD8",
        "main_bg":      "#FFFDF9",
        "text_primary": "#3D2B1F",
        "text_muted":   "#8C6B55",
        "accent":       "#C0533A",
        "accent_light": "#FDEEE9",
        "header_bg":    "#3D2B1F",
        "header_text":  "#FDF6EC",
        "sidebar_text": "#3D2B1F",
        "sidebar_muted":"#8C6B55",
        "divider":      "#E8D5C0",
    },
    {
        "label": "Arctic Blue & Steel",
        "page_bg":      "#F0F4F8",
        "sidebar_bg":   "#DBEAFE",
        "main_bg":      "#FFFFFF",
        "text_primary": "#1E3A5F",
        "text_muted":   "#64748B",
        "accent":       "#1D4ED8",
        "accent_light": "#EFF6FF",
        "header_bg":    "#1D4ED8",
        "header_text":  "#FFFFFF",
        "sidebar_text": "#1E3A5F",
        "sidebar_muted":"#64748B",
        "divider":      "#BFDBFE",
    },
    {
        "label": "Graphite & Coral",
        "page_bg":      "#F7F7F7",
        "sidebar_bg":   "#2D2D2D",
        "main_bg":      "#FFFFFF",
        "text_primary": "#2D2D2D",
        "text_muted":   "#737373",
        "accent":       "#FF6B6B",
        "accent_light": "#FFF0F0",
        "header_bg":    "#2D2D2D",
        "header_text":  "#FFFFFF",
        "sidebar_text": "#F5F5F5",
        "sidebar_muted":"#A3A3A3",
        "divider":      "#E5E5E5",
    },
    {
        "label": "Lavender & Charcoal",
        "page_bg":      "#FAF9FF",
        "sidebar_bg":   "#EDE9FE",
        "main_bg":      "#FFFFFF",
        "text_primary": "#1C1917",
        "text_muted":   "#78716C",
        "accent":       "#7C3AED",
        "accent_light": "#F5F3FF",
        "header_bg":    "#4C1D95",
        "header_text":  "#FFFFFF",
        "sidebar_text": "#1C1917",
        "sidebar_muted":"#78716C",
        "divider":      "#DDD6FE",
    },
]

_COMPONENTS = [
    {
        "label": "Sharp & Minimal",
        "skill_style": (
            "inline-block; padding: 2px 8px; border: 1px solid {accent}; "
            "border-radius: 0; font-size: 9px; color: {accent}; margin: 2px 2px 2px 0;"
        ),
        "divider_style": "border: none; border-top: 1px solid {divider}; margin: 12px 0;",
        "experience_style": (
            "Flat entries, no timeline decoration. Role title bold in accent colour, "
            "company + dates on the same line below in muted colour, dates right-aligned."
        ),
        "heading_decoration": (
            "Plain text, all-caps, letter-spaced, no background, accent colour, "
            "no border — rely on spacing above the heading only."
        ),
    },
    {
        "label": "Rounded Pills",
        "skill_style": (
            "inline-block; padding: 3px 10px; background: {accent_light}; "
            "border-radius: 20px; font-size: 9px; color: {accent}; "
            "margin: 2px 2px 2px 0; font-weight: 500;"
        ),
        "divider_style": "border: none; border-top: 2px solid {accent}; margin: 10px 0;",
        "experience_style": (
            "Clean entries. Role title in accent colour bold, company on next line italic muted, "
            "dates right-aligned using flex justify-between on the title row."
        ),
        "heading_decoration": (
            "Left border 3px solid accent, 8px left-padding, background accent_light, "
            "small bold all-caps text — a coloured side-bar block per heading."
        ),
    },
    {
        "label": "Vertical Timeline",
        "skill_style": (
            "inline-block; padding: 2px 6px; background: {accent_light}; "
            "border-radius: 3px; font-size: 9px; color: {text_primary}; margin: 2px 2px 2px 0;"
        ),
        "divider_style": "border: none; border-top: 1px dashed {divider}; margin: 12px 0;",
        "experience_style": (
            "Each experience entry has a continuous 2px solid left border in accent colour "
            "with 14px left padding — creating a vertical timeline effect. "
            "Date appears in small monospace muted font at top-right of each entry block."
        ),
        "heading_decoration": (
            "Bold all-caps, underlined with a 2px solid accent-colour bottom border "
            "only on the heading text itself, not full-width."
        ),
    },
    {
        "label": "Dashed Borders",
        "skill_style": (
            "inline-block; padding: 2px 6px; border: 1px dashed {text_muted}; "
            "border-radius: 2px; font-size: 9px; color: {text_muted}; margin: 2px 2px 2px 0;"
        ),
        "divider_style": "border: none; border-top: 1px dashed {divider}; margin: 14px 0;",
        "experience_style": (
            "Minimal entries — bold role title, italic muted company name, "
            "date right-aligned in a small monospace font. "
            "Tight line-height, generous margin between entries."
        ),
        "heading_decoration": (
            "Italic small-caps in accent colour, with a small bullet · prefix, "
            "no background, no border — elegant and understated."
        ),
    },
    {
        "label": "Filled Badges",
        "skill_style": (
            "inline-block; padding: 3px 8px; background: {accent}; "
            "border-radius: 3px; font-size: 8.5px; color: white; "
            "margin: 2px 2px 2px 0; font-weight: 600; "
            "text-transform: uppercase; letter-spacing: 0.5px;"
        ),
        "divider_style": "border: none; border-top: 3px double {accent}; margin: 10px 0;",
        "experience_style": (
            "Each entry: accent-coloured role title, company on next line, "
            "achievements as a tight bulleted list directly below. "
            "No separate date column — date is inline after company name in parentheses, muted."
        ),
        "heading_decoration": (
            "Full-width accent_light background block, 4px top/bottom and 8px left/right padding, "
            "small bold all-caps text in accent colour — a coloured band for each section heading."
        ),
    },
    {
        "label": "Underline Accent",
        "skill_style": (
            "inline-block; padding: 1px 4px; "
            "border-bottom: 1.5px solid {accent}; "
            "font-size: 10px; color: {text_primary}; margin: 2px 4px 2px 0;"
        ),
        "divider_style": "border: none; border-bottom: 0.5px solid {divider}; margin: 16px 0;",
        "experience_style": (
            "Generous spacing. Role title bold in accent colour. "
            "Company and date on same line: company normal weight, date in small muted monospace at far right. "
            "Bullet points with standard body font, 10px, normal weight."
        ),
        "heading_decoration": (
            "Large font-weight 300, no letter-spacing, accent bottom-border only (not underline, "
            "use border-bottom with padding-bottom), no transform — clean and open."
        ),
    },
]


def _generate_random_design_spec() -> dict:
    """
    Procedurally build a unique design spec by sampling one item from each
    of the four independent design axes.

    Returns {"name": display_label, "description": prompt_instructions}.
    The name is injected into the UI badge; the description is injected
    verbatim into the HTML-generation prompt.
    """
    structure  = random.choice(_STRUCTURES)
    typography = random.choice(_TYPOGRAPHY)
    palette    = random.choice(_PALETTES)
    component  = random.choice(_COMPONENTS)

    name = (
        f"{typography['label']} "
        f"• {palette['label']} "
        f"• {structure['label']}"
    )

    # Substitute palette + typography tokens into component style strings
    def sub(s: str) -> str:
        return (
            s.replace("{accent}",       palette["accent"])
             .replace("{accent_light}", palette["accent_light"])
             .replace("{divider}",      palette["divider"])
             .replace("{text_primary}", palette["text_primary"])
             .replace("{text_muted}",   palette["text_muted"])
             .replace("{heading_size}", typography["heading_size"])
        )

    description = f"""
STRUCTURE:
{structure['description']}

TYPOGRAPHY:
  Google Fonts: {typography['import']}
  Heading font-family: {typography['heading_font']}
  Body font-family:    {typography['body_font']}
  Candidate name font-size: {typography['name_size']}
  Section heading: font-size {typography['heading_size']}, font-weight {typography['heading_weight']}, text-transform {typography['heading_transform']}, letter-spacing {typography['heading_spacing']}
  Body text font-size: {typography['body_size']}

COLOUR PALETTE — {palette['label']}:
  Page background:        {palette['page_bg']}
  Sidebar background:     {palette['sidebar_bg']}
  Main/content background:{palette['main_bg']}
  Primary text:           {palette['text_primary']}
  Muted / secondary text: {palette['text_muted']}
  Accent (headings, name):{palette['accent']}
  Accent light (tag bgs): {palette['accent_light']}
  Header band background: {palette['header_bg']}
  Header band text:       {palette['header_text']}
  Sidebar text colour:    {palette['sidebar_text']}
  Sidebar muted text:     {palette['sidebar_muted']}
  Divider / rule colour:  {palette['divider']}

COMPONENT STYLING — {component['label']}:
  Skill tag inline style:   "{sub(component['skill_style'])}"
  Section divider style:    "{sub(component['divider_style'])}"
  Experience entry layout:  {sub(component['experience_style'])}
  Section heading decoration: {sub(component['heading_decoration'])}

CRITICAL: implement every colour, font, and component style exactly as specified.
Do not substitute or approximate — use the exact hex values and CSS properties above.
""".strip()

    return {"name": name, "description": description}


# ── Profile serialisation ──────────────────────────────────────────────────────

def _build_profile_text(user_id: str) -> str:
    """Convert the user's real profile (get_profile(user_id)) into a structured text block for prompts."""
    p    = get_profile(user_id)
    name = p["personal"]["name"] or "(name not yet provided — do not invent one)"
    lines: list[str] = [f"CANDIDATE: {name}\n"]

    if goals := (p.get("career_goals") or {}).get("target_roles"):
        lines.append(f"TARGET ROLE(S) (candidate's own stated goal): {', '.join(goals)}")

    if self_summary := (p.get("summary") or "").strip():
        lines.append(f"CANDIDATE'S OWN SELF-SUMMARY (mine for voice/facts, do not copy verbatim):\n  {self_summary}\n")

    lines.append("EDUCATION:")
    for edu in p.get("education", []):
        if "degree" in edu:
            lines.append(
                f"  • {edu['degree']} — {edu.get('school', '')} [{edu.get('status', '')}]"
            )
            if note := edu.get("resilience_note"):
                lines.append(f"    Context: {note}")
        if "certification" in edu:
            lines.append(
                f"  • Certification: {edu['certification']} from {edu.get('provider', '')} "
                f"({edu.get('details', '')})"
            )

    lines.append("\nEXPERIENCE:")
    for exp in p.get("experience", []):
        company = exp.get("company") or exp.get("unit", "")
        period  = exp.get("period", "")
        if "roles" in exp:
            for r in exp["roles"]:
                lines.append(f"  • {r['title']} at {company} ({r['period']})")
                lines.append(f"    {r['details']}")
        else:
            role    = exp.get("role", "")
            details = exp.get("details", "")
            lines.append(f"  • {role} at {company} ({period})")
            if details:
                lines.append(f"    {details}")

    vol = p.get("volunteering", {})
    if vol:
        lines.append("\nVOLUNTEERING:")
        if isinstance(vol, dict):
            lines.append(
                f"  • {vol.get('role', '')} at {vol.get('organization', '')} "
                f"({vol.get('duration', '')}): {vol.get('description', '')}"
            )
        else:
            lines.append(f"  • {vol}")

    skills = p.get("skills", [])
    if skills:
        lines.append(f"\nSKILLS: {', '.join(skills)}")

    narratives = p.get("key_narratives", {})
    if narratives:
        lines.append("\nKEY CAREER NARRATIVES (use these to write impact bullets):")
        for key, n in narratives.items():
            lines.append(f"  [{key}] {n.get('headline', '')}")
            for ev in n.get("evidence", []):
                lines.append(f"    - {ev}")

    return "\n".join(lines)


# ── Agent ──────────────────────────────────────────────────────────────────────

class ResumeAgent:
    def __init__(self) -> None:
        self._client = anthropic.AsyncAnthropic()

    async def generate(
        self,
        job: JobMatch,
        user_id: str,
        supplemental_answers: Optional[dict[str, str]] = None,
        reference_bytes: Optional[bytes] = None,
        reference_mime: Optional[str] = None,
    ) -> tuple[str, list[dict], str]:
        """
        Returns (html, missing_data_requests, layout_variant_name).

        missing_data_requests: list of {"id", "question", "context"} dicts.
        layout_variant_name:   the layout used (for display in the UI).
        """
        # 1. Optional: extract style cues from the reference file (must run first)
        style_override: str = ""
        if reference_bytes and reference_mime:
            logger.info("[resume] Analysing reference file (%s, %d bytes)…", reference_mime, len(reference_bytes))
            style_override = await self._analyse_reference_file(reference_bytes, reference_mime)
            logger.info("[resume] Style extracted: %s…", style_override[:120])

        # Build design spec — use vision-extracted style if available, otherwise procedurally generated
        if style_override:
            layout = {"name": "Reference Style", "description": style_override}
        else:
            layout = _generate_random_design_spec()

        logger.info("[resume] Layout: %s", layout["name"])

        # 2. Gap analysis + HTML generation in parallel
        profile_text = _build_profile_text(user_id)
        gaps_coro    = self._analyse_gaps(job, supplemental_answers or {}, profile_text)
        html_coro    = self._generate_html(job, supplemental_answers or {}, profile_text, layout)

        missing_data, html = await asyncio.gather(gaps_coro, html_coro)

        logger.info(
            "[resume] Done — %d gap question(s), HTML %d chars",
            len(missing_data), len(html),
        )
        return html, missing_data, layout["name"]

    # ── Private helpers ────────────────────────────────────────────────────────

    _STYLE_ANALYSIS_PROMPT = (
        "Analyse this resume and describe its visual layout precisely so I can replicate it in HTML/CSS. "
        "Cover: (1) overall column structure (1-col, 2-col, sidebar — include approximate proportions), "
        "(2) header placement and style (centred, left-aligned, coloured band, etc.), "
        "(3) typography — serif vs sans-serif, approximate font sizes for name/headings/body, "
        "(4) colour scheme — infer hex codes where possible, otherwise describe as 'dark navy', 'warm sand', etc., "
        "(5) section heading decoration (borders, backgrounds, small-caps, all-caps, underlines), "
        "(6) skill/tag style (plain text, chips, badges, comma-separated), "
        "(7) experience entry layout (timeline, flat, left-border, etc.), "
        "(8) section order top-to-bottom. "
        "Return a single structured paragraph of CSS/HTML layout instructions a developer can follow directly. "
        "Be specific about pixel sizes, column widths, and colour values."
    )

    async def _analyse_reference_file(self, file_bytes: bytes, mime: str) -> str:
        """Dispatch to the correct analyser based on MIME type."""
        if mime.startswith("image/"):
            return await self._analyse_image(file_bytes, mime)
        if mime == "application/pdf":
            return await self._analyse_pdf(file_bytes)
        if "wordprocessingml" in mime:
            return await self._analyse_docx(file_bytes)
        # Fallback: treat as image
        return await self._analyse_image(file_bytes, mime)

    async def _analyse_image(self, img_bytes: bytes, media_type: str) -> str:
        """Send an image to Claude Vision for layout analysis."""
        b64 = base64.standard_b64encode(img_bytes).decode()
        result = await call_llm(
            messages=[{
                "role": "user",
                "content": [
                    {"type": "image", "source": {"type": "base64", "media_type": media_type, "data": b64}},
                    {"type": "text", "text": self._STYLE_ANALYSIS_PROMPT},
                ],
            }],
            model=_MODEL,
            max_tokens=600,
            purpose="resume_analyse_image",
        )
        return result.text.strip()

    async def _analyse_pdf(self, pdf_bytes: bytes) -> str:
        """
        Analyse a PDF reference resume.

        Attempt 1 — Anthropic native PDF support (beta, no rendering needed).
        Attempt 2 — PyMuPDF: render page 1 as a PNG and send via Vision.
        """
        b64 = base64.standard_b64encode(pdf_bytes).decode()
        try:
            msg = await self._client.beta.messages.create(
                model=_MODEL,
                max_tokens=600,
                betas=["pdfs-2024-09-25"],
                messages=[{
                    "role": "user",
                    "content": [
                        {
                            "type": "document",
                            "source": {"type": "base64", "media_type": "application/pdf", "data": b64},
                        },
                        {"type": "text", "text": self._STYLE_ANALYSIS_PROMPT},
                    ],
                }],
            )
            return msg.content[0].text.strip()
        except Exception as exc:
            logger.warning("[resume] Native PDF analysis failed (%s), falling back to PyMuPDF", exc)

        # Fallback: render first page with PyMuPDF
        import fitz  # PyMuPDF
        doc = fitz.open(stream=pdf_bytes, filetype="pdf")
        page = doc[0]
        # 2× zoom gives ~150 dpi — good enough for layout analysis without huge payload
        pix = page.get_pixmap(matrix=fitz.Matrix(2.0, 2.0))
        png_bytes = pix.tobytes("png")
        doc.close()
        logger.info("[resume] PyMuPDF rendered page 1 as PNG (%d bytes)", len(png_bytes))
        return await self._analyse_image(png_bytes, "image/png")

    async def _analyse_docx(self, docx_bytes: bytes) -> str:
        """
        Extract text and heading structure from a .docx file, then ask Claude
        to infer the visual layout from the document structure.
        """
        from docx import Document
        from docx.oxml.ns import qn

        doc = Document(io.BytesIO(docx_bytes))

        # Build a structured representation: heading levels + body snippets
        lines: list[str] = []
        for para in doc.paragraphs:
            style = para.style.name if para.style else ""
            text  = para.text.strip()
            if not text:
                continue
            if style.startswith("Heading"):
                level = style.replace("Heading", "").strip() or "1"
                lines.append(f"[H{level}] {text}")
            else:
                # Body paragraph — keep but truncate long ones
                lines.append(text[:120] + ("…" if len(text) > 120 else ""))

        # Limit to ~100 lines to stay within token budget
        structure_text = "\n".join(lines[:100])

        # Also capture table content (some resumes use tables for layout)
        table_snippets: list[str] = []
        for table in doc.tables[:4]:
            row_texts = [
                " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                for row in table.rows
            ]
            table_snippets.append("\n".join(r for r in row_texts if r))

        table_block = ""
        if table_snippets:
            table_block = "\n\nTABLES FOUND (layout tables are common in Word resumes):\n" + "\n---\n".join(table_snippets[:3])

        prompt = (
            "Below is text extracted from a Word (.docx) resume. "
            "The headings (marked [H1], [H2]) reveal the document structure; "
            "table data (if present) often indicates a multi-column layout. "
            "Based on the structure, section order, and any table layout, describe the visual design "
            "of this resume and provide CSS/HTML layout instructions to replicate it.\n\n"
            "DOCUMENT STRUCTURE:\n"
            + structure_text
            + table_block
            + "\n\n"
            + self._STYLE_ANALYSIS_PROMPT
        )

        result = await call_llm(
            messages=[{"role": "user", "content": prompt}],
            model=_MODEL,
            max_tokens=600,
            purpose="resume_analyse_docx",
        )
        return result.text.strip()

    async def _analyse_gaps(
        self,
        job: JobMatch,
        supplemental_answers: dict[str, str],
        profile_text: str,
    ) -> list[dict]:
        """Identify what the JD requires that isn't evidenced in the profile."""
        already_answered = (
            "\n".join(f"- {k}: {v}" for k, v in supplemental_answers.items())
            if supplemental_answers
            else "None."
        )

        job_context = (
            f"Job Title: {job.title}\n"
            f"Company: {job.company}\n"
            f"Location: {job.location}\n"
            f"Match reasons: {', '.join(r.label for r in job.reasons)}\n"
            f"Investigation points (known gaps/questions): {'; '.join(job.investigation_points)}\n"
            f"Trajectory note: {job.trajectory_alignment}"
        )

        prompt = (
            "You are a professional resume analyst.\n\n"
            "JOB:\n" + job_context + "\n\n"
            "CANDIDATE PROFILE:\n" + profile_text + "\n\n"
            "ALREADY ANSWERED:\n" + already_answered + "\n\n"
            "Identify up to 5 pieces of information that would STRENGTHEN this resume for this specific job "
            "but are NOT clearly evidenced in the profile and NOT already answered above. "
            "Only ask about genuinely missing concrete data (numbers, technologies, dates, certifications). "
            "Do NOT ask about things the profile already covers. "
            "If the profile fully covers the role, return an empty array.\n\n"
            'Return ONLY a JSON array. Each element: {"id": "snake_case_id", "question": "direct question for the candidate", "context": "why this matters for this specific job"}. '
            "No markdown, no explanation, no code fences — raw JSON only."
        )

        result = await call_llm(
            messages=[{"role": "user", "content": prompt}],
            model=_MODEL,
            max_tokens=600,
            purpose="resume_analyse_gaps",
        )
        raw = result.text.strip()

        try:
            data = json.loads(raw)
            if isinstance(data, list):
                return data[:5]
        except (json.JSONDecodeError, TypeError):
            logger.warning("[resume] Gap analysis returned non-JSON: %s", raw[:200])
        return []

    async def _generate_html(
        self,
        job: JobMatch,
        supplemental_answers: dict[str, str],
        profile_text: str,
        layout: dict,
    ) -> str:
        """Generate a complete, self-contained HTML resume document."""
        supplemental_block = (
            "\n".join(f"  {k}: {v}" for k, v in supplemental_answers.items())
            if supplemental_answers
            else "  None — use only the profile data above."
        )

        reasons_text = ", ".join(r.label for r in job.reasons[:6])
        investigation = "; ".join(job.investigation_points[:4]) if job.investigation_points else "N/A"

        prompt = f"""You are a world-class resume designer generating a complete HTML document.

LAYOUT TO IMPLEMENT:
{layout['description']}

CANDIDATE PROFILE:
{profile_text}

SUPPLEMENTAL INFORMATION FROM CANDIDATE:
{supplemental_block}

JOB TARGET:
  Title: {job.title}
  Company: {job.company}
  Location: {job.location}
  Matched signals: {reasons_text}
  Areas to emphasise: {investigation}
  Trajectory: {job.trajectory_alignment}

TAILORING INSTRUCTIONS:
- Professional Summary — write AT LEAST 4 detailed lines (not 4 short fragments; 4 full, substantive lines of running text or up to 4 tight sentences). It must read like a senior/executive-caliber summary, not a junior objective statement. Cover, in this order:
    1. Years of experience + functional domain, framed against "{job.title}" at "{job.company}".
    2. The 1-2 most quantified, relevant achievements from the profile (real numbers/scope only — never invented).
    3. Core capabilities that directly mirror the job's matched signals ({reasons_text}), stated as evidenced strengths, not a skills list restated.
    4. A trajectory/scope signal — leadership scope, cross-functional reach, or specialized depth — that shows why this candidate's level fits the role. Do NOT penalize a career pivot or treat greater seniority/experience than the role requires as a negative; frame it as range and readiness, never as a mismatch.
  Never use unsupported filler ("results-driven professional", "proven track record", "team player") without a concrete fact backing it up in the same sentence.
- Bullet points — 2 to 4 per role, outcome-first using the XYZ pattern: accomplished [X], measured by [Y — a real number/scope from the profile], by doing [Z]. Only quantify with numbers/scope actually present in the profile or supplemental answers; if no number exists, express scale qualitatively (e.g. "cut deployment time from days to hours") rather than fabricating a metric.
- Reorder bullet points within each role so the most relevant achievements to this job appear first.
- Mirror vocabulary from the job title and matched signals naturally in bullet text — do not keyword-stuff.
- Build the Skills/Core Competencies section by prioritizing profile skills that overlap with this job's matched signals and areas to emphasise ({investigation}); never list a skill not evidenced somewhere in the profile.
- Incorporate any supplemental answers as concrete claims in the relevant role or skills section.

ATS COMPLIANCE RULES (non-negotiable):
- Emit a complete HTML5 document (<!DOCTYPE html><html><head>…</head><body>…</body></html>).
- All CSS in a <style> block in <head>. You MAY use one Google Fonts @import.
- ZERO external CSS files or JS files. ZERO tables for any layout purpose. ZERO images, icons, or emoji as bullet markers — use a plain CSS list-style or a simple text dash/bullet character only.
- CRITICAL — DOM reading order is independent of visual layout: regardless of the visual structure above (including any sidebar/multi-column look), the underlying HTML source order MUST be single-file linear reading order — name → contact → summary → experience (most recent first) → education → skills — exactly as a human would scan it top-to-bottom. Use CSS (flexbox/grid/floats), never HTML tables or absolute positioning that would let a visual sidebar's DOM position diverge from this reading order. ATS parsers read raw DOM/text order, not visual position, and multi-column HTML is the single largest cause of ATS parsing failure when source order doesn't match visual order.
- Use standard, ATS-recognized section headings only: "Summary" (or "Professional Summary"), "Experience", "Education", "Skills" — never creative alternatives like "What I've Done".
- Use semantic tags: <h1> for candidate name only, <h2> for section headings, <h3> for job titles.
- Contact line must use literal placeholders: [Email] · [Phone] · LinkedIn · {job.location or 'Location'}.
- Dates must use "MMM YYYY" format (e.g. "Jan 2022 – Mar 2024"), never seasons or apostrophe-shortened years.
- Every bullet must start with a strong past-tense action verb.
- Do NOT invent facts, numbers, or skills not present in the profile or supplemental answers.
- The document must render correctly in an iframe at standard A4/letter proportions.

OUTPUT: The complete HTML document only. No markdown. No explanation. No code fences."""

        result = await call_llm(
            messages=[{"role": "user", "content": prompt}],
            model=_MODEL,
            max_tokens=_MAX_TOKENS,
            purpose="resume_generate_html",
        )
        html = result.text.strip()

        # Strip accidental markdown fences if model adds them
        if html.startswith("```"):
            html = html.split("\n", 1)[1] if "\n" in html else html[3:]
        if html.endswith("```"):
            html = html.rsplit("```", 1)[0].rstrip()

        return html
